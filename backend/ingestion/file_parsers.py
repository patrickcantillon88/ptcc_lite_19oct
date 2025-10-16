"""
File parsers for different document types
"""

import os
import re
import json
from typing import Dict, List, Any, Optional
from pathlib import Path

import pandas as pd
import pdfplumber
import docx
from datetime import datetime

from ..core.logging_config import get_logger

logger = get_logger("file_parsers")


class BaseParser:
    """Base class for file parsers"""
    
    def __init__(self):
        self.supported_extensions = []
    
    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the given file"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse the file and return structured data"""
        raise NotImplementedError("Subclasses must implement parse method")
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract basic metadata from file"""
        file_stat = os.stat(file_path)
        return {
            "file_name": os.path.basename(file_path),
            "file_path": file_path,
            "file_size": file_stat.st_size,
            "created_date": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            "modified_date": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
            "file_type": Path(file_path).suffix.lower()
        }


class ExcelParser(BaseParser):
    """Parser for Excel files (.xlsx, .xls)"""
    
    def __init__(self):
        self.supported_extensions = ['.xlsx', '.xls']
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse Excel file and extract data from all sheets"""
        try:
            metadata = self.extract_metadata(file_path)
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert DataFrame to records
                    records = df.to_dict('records')
                    
                    # Clean up NaN values
                    cleaned_records = []
                    for record in records:
                        cleaned_record = {}
                        for key, value in record.items():
                            if pd.isna(value):
                                cleaned_record[key] = None
                            elif isinstance(value, str):
                                cleaned_record[key] = value.strip()
                            else:
                                cleaned_record[key] = value
                        cleaned_records.append(cleaned_record)
                    
                    sheets_data[sheet_name] = {
                        "columns": list(df.columns),
                        "data": cleaned_records,
                        "row_count": len(cleaned_records),
                        "column_count": len(df.columns)
                    }
                    
                except Exception as e:
                    logger.warning(f"Error reading sheet '{sheet_name}' in {file_path}: {e}")
                    sheets_data[sheet_name] = {
                        "error": str(e),
                        "columns": [],
                        "data": [],
                        "row_count": 0,
                        "column_count": 0
                    }
            
            return {
                **metadata,
                "file_type": "excel",
                "sheets": sheets_data,
                "sheet_count": len(sheets_data)
            }
            
        except Exception as e:
            logger.error(f"Error parsing Excel file {file_path}: {e}")
            raise ValueError(f"Failed to parse Excel file: {e}")
    
    def detect_file_type(self, file_path: str) -> Optional[str]:
        """Detect the type of Excel file based on content"""
        try:
            data = self.parse(file_path)
            
            # Check for class list patterns
            for sheet_name, sheet_data in data["sheets"].items():
                if sheet_data["row_count"] == 0:
                    continue
                
                columns = [col.lower() for col in sheet_data["columns"]]
                
                # Class list indicators
                if any(col in columns for col in ["name", "student", "pupil", "class", "form"]):
                    return "class_list"
                
                # Assessment indicators
                if any(col in columns for col in ["score", "grade", "mark", "percentage", "assessment"]):
                    return "assessment"
                
                # Timetable indicators
                if any(col in columns for col in ["period", "time", "day", "subject", "room"]):
                    return "timetable"
            
            return "unknown"
            
        except Exception:
            return "unknown"


class PDFParser(BaseParser):
    """Parser for PDF files"""
    
    def __init__(self):
        self.supported_extensions = ['.pdf']
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF file and extract text content with enhanced table extraction"""
        try:
            metadata = self.extract_metadata(file_path)

            with pdfplumber.open(file_path) as pdf:
                pages_data = []
                full_text = ""
                all_tables = []

                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text() or ""
                        full_text += page_text + "\n"

                        # Enhanced table extraction with better handling
                        tables = []
                        try:
                            # Extract tables with explicit table settings for better detection
                            page_tables = page.extract_tables(table_settings={
                                "vertical_strategy": "lines",
                                "horizontal_strategy": "lines",
                                "explicit_vertical_lines": [],
                                "explicit_horizontal_lines": [],
                                "snap_tolerance": 3,
                                "join_tolerance": 3,
                                "edge_min_length": 3,
                                "min_words_vertical": 3,
                                "min_words_horizontal": 1,
                                "keep_blank_chars": False,
                                "text_tolerance": 1,
                                "text_x_tolerance": None,
                                "text_y_tolerance": None,
                                "intersection_tolerance": 3,
                                "intersection_x_tolerance": None,
                                "intersection_y_tolerance": None,
                            })

                            for table_idx, table in enumerate(page_tables):
                                # Clean up table data and handle merged cells
                                cleaned_table = []
                                if table:  # Ensure table is not empty
                                    for row in table:
                                        cleaned_row = []
                                        for cell in row:
                                            if cell is None:
                                                cleaned_row.append("")
                                            else:
                                                # Clean up cell content
                                                cell_text = str(cell).strip()
                                                # Remove excessive whitespace and normalize
                                                cell_text = ' '.join(cell_text.split())
                                                cleaned_row.append(cell_text)
                                        cleaned_table.append(cleaned_row)

                                    # Only add non-empty tables
                                    if cleaned_table and any(any(cell for cell in row) for row in cleaned_table):
                                        tables.append(cleaned_table)
                                        all_tables.append({
                                            "page": page_num,
                                            "table_index": table_idx,
                                            "data": cleaned_table,
                                            "row_count": len(cleaned_table),
                                            "col_count": len(cleaned_table[0]) if cleaned_table else 0
                                        })

                        except Exception as e:
                            logger.debug(f"Could not extract tables from page {page_num}: {e}")

                        pages_data.append({
                            "page_number": page_num,
                            "text": page_text,
                            "tables": tables,
                            "table_count": len(tables)
                        })

                    except Exception as e:
                        logger.warning(f"Error processing page {page_num} in {file_path}: {e}")
                        pages_data.append({
                            "page_number": page_num,
                            "text": "",
                            "tables": [],
                            "table_count": 0,
                            "error": str(e)
                        })

                return {
                    **metadata,
                    "file_type": "pdf",
                    "pages": pages_data,
                    "page_count": len(pages_data),
                    "full_text": full_text.strip(),
                    "all_tables": all_tables,
                    "total_tables": len(all_tables)
                }

        except Exception as e:
            logger.error(f"Error parsing PDF file {file_path}: {e}")
            raise ValueError(f"Failed to parse PDF file: {e}")
    
    def detect_file_type(self, file_path: str) -> Optional[str]:
        """Detect the type of PDF based on content"""
        try:
            data = self.parse(file_path)
            text = data["full_text"].lower()
            
            # Class list indicators
            if any(keyword in text for keyword in ["class list", "student list", "roll", "register"]):
                return "class_list"
            
            # Assessment indicators
            if any(keyword in text for keyword in ["assessment", "exam", "test", "score", "grade"]):
                return "assessment"
            
            # Meeting minutes indicators
            if any(keyword in text for keyword in ["meeting", "minutes", "agenda", "attendance"]):
                return "meeting_minutes"
            
            # Report indicators
            if any(keyword in text for keyword in ["report", "progress", "evaluation", "performance"]):
                return "report"
            
            return "document"
            
        except Exception:
            return "unknown"


class WordParser(BaseParser):
    """Parser for Word documents (.docx)"""
    
    def __init__(self):
        self.supported_extensions = ['.docx']
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse Word document and extract text content"""
        try:
            metadata = self.extract_metadata(file_path)
            
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text.strip(),
                        "style": para.style.name if para.style else "Normal"
                    })
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Combine all text
            full_text = "\n".join(para["text"] for para in paragraphs)
            
            return {
                **metadata,
                "file_type": "word",
                "paragraphs": paragraphs,
                "paragraph_count": len(paragraphs),
                "tables": tables,
                "table_count": len(tables),
                "full_text": full_text
            }
            
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {e}")
            raise ValueError(f"Failed to parse Word document: {e}")
    
    def detect_file_type(self, file_path: str) -> Optional[str]:
        """Detect the type of Word document based on content"""
        try:
            data = self.parse(file_path)
            text = data["full_text"].lower()
            
            # Class list indicators
            if any(keyword in text for keyword in ["class list", "student list", "roll", "register"]):
                return "class_list"
            
            # Assessment indicators
            if any(keyword in text for keyword in ["assessment", "exam", "test", "score", "grade"]):
                return "assessment"
            
            # Meeting minutes indicators
            if any(keyword in text for keyword in ["meeting", "minutes", "agenda", "attendance"]):
                return "meeting_minutes"
            
            # Report indicators
            if any(keyword in text for keyword in ["report", "progress", "evaluation", "performance"]):
                return "report"
            
            # Communication indicators
            if any(keyword in text for keyword in ["newsletter", "notice", "announcement", "circular"]):
                return "communication"
            
            return "document"
            
        except Exception:
            return "unknown"


class JSONParser(BaseParser):
    """Parser for JSON files containing RAG data"""

    def __init__(self):
        self.supported_extensions = ['.json']

    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse JSON file containing RAG data"""
        try:
            metadata = self.extract_metadata(file_path)

            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)

            # Validate structure - should be list of objects with content and metadata
            if not isinstance(json_data, list):
                raise ValueError("JSON file must contain a list of RAG data objects")

            # Validate first few items
            for i, item in enumerate(json_data[:5]):
                if not isinstance(item, dict):
                    raise ValueError(f"Item {i} is not a dictionary")
                if 'content' not in item:
                    raise ValueError(f"Item {i} missing 'content' field")
                if 'metadata' not in item:
                    raise ValueError(f"Item {i} missing 'metadata' field")

            return {
                **metadata,
                "file_type": "json",
                "rag_data": json_data,
                "item_count": len(json_data)
            }

        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON file: {e}")
        except Exception as e:
            raise ValueError(f"Failed to parse JSON file: {e}")

    def detect_file_type(self, file_path: str) -> Optional[str]:
        """Detect if this is RAG data JSON"""
        try:
            data = self.parse(file_path)
            # Check if it looks like RAG data
            if data.get("item_count", 0) > 0:
                first_item = data["rag_data"][0]
                if "content" in first_item and "metadata" in first_item:
                    return "rag_data"
            return "json"
        except Exception:
            return "unknown"


class FileParserFactory:
    """Factory class to get appropriate parser for a file"""

    def __init__(self):
        self.parsers = [
            ExcelParser(),
            PDFParser(),
            WordParser(),
            JSONParser()
        ]
    
    def get_parser(self, file_path: str) -> Optional[BaseParser]:
        """Get the appropriate parser for the given file"""
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser
        return None
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """Parse a file using the appropriate parser"""
        parser = self.get_parser(file_path)
        if not parser:
            raise ValueError(f"No parser found for file: {file_path}")
        
        return parser.parse(file_path)
    
    def detect_file_type(self, file_path: str) -> str:
        """Detect the type of file based on content"""
        parser = self.get_parser(file_path)
        if not parser:
            return "unknown"
        
        # Try to detect specific file type
        if hasattr(parser, 'detect_file_type'):
            return parser.detect_file_type(file_path)
        
        # Fallback to generic file type
        return parser.__class__.__name__.replace("Parser", "").lower()